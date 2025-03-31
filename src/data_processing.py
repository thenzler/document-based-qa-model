"""
Data Processing Module for Document-based QA System

This module handles all the document processing, including:
- Loading documents from various formats (txt, pdf, docx, etc.)
- Preprocessing text (tokenization, cleaning, etc.)
- Creating document embeddings
- Indexing documents for fast retrieval
"""

import os
import re
import glob
import pandas as pd
import numpy as np
from typing import List, Dict, Union, Tuple, Optional, Set
from pathlib import Path
import json
import logging

# For text processing
from transformers import AutoTokenizer
from sentence_transformers import SentenceTransformer
import spacy

# For document indexing
import faiss
from langchain.text_splitter import RecursiveCharacterTextSplitter

# For document parsing
import fitz  # PyMuPDF for PDF
import docx  # python-docx for DOCX
import pdfplumber
from bs4 import BeautifulSoup
import nltk
from tqdm import tqdm

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

try:
    nltk.data.find('tokenizers/punkt')
except:
    nltk.download('punkt')


class DocumentProcessor:
    """Handles document loading, preprocessing, and indexing"""
    
    def __init__(self, model_name: str = "sentence-transformers/all-mpnet-base-v2"):
        """
        Initialize the document processor.
        
        Args:
            model_name: Name of the sentence transformer model to use for embeddings
        """
        self.embedding_model = SentenceTransformer(model_name)
        try:
            self.nlp = spacy.load("de_core_news_md")
        except OSError:
            # Download if not available
            os.system("python -m spacy download de_core_news_md")
            self.nlp = spacy.load("de_core_news_md")
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        self.document_store = []
        self.document_embeddings = None
        self.index = None
        self.section_markers = [
            r'^\d+\.\s+\w+',  # Numbered sections like "1. Introduction"
            r'^#+\s+\w+',     # Markdown headings
            r'^\w+:',         # Section headers ending with colon
        ]
    
    def load_documents(self, docs_dir: Union[str, Path]) -> List[Dict]:
        """
        Load all documents from a directory.
        
        Args:
            docs_dir: Directory containing document files
            
        Returns:
            List of document dictionaries with text and metadata
        """
        docs_dir = Path(docs_dir)
        
        # List all supported file types in the directory
        text_files = list(docs_dir.glob("**/*.txt"))
        markdown_files = list(docs_dir.glob("**/*.md"))
        pdf_files = list(docs_dir.glob("**/*.pdf"))
        docx_files = list(docs_dir.glob("**/*.docx"))
        html_files = list(docs_dir.glob("**/*.html"))
        
        all_files = text_files + markdown_files + pdf_files + docx_files + html_files
        documents = []
        
        for file_path in tqdm(all_files, desc="Loading documents"):
            try:
                # Handle different file types
                if file_path.suffix.lower() == '.txt' or file_path.suffix.lower() == '.md':
                    with open(file_path, "r", encoding="utf-8") as f:
                        text = f.read()
                
                elif file_path.suffix.lower() == '.pdf':
                    text = self._extract_text_from_pdf(file_path)
                
                elif file_path.suffix.lower() == '.docx':
                    text = self._extract_text_from_docx(file_path)
                
                elif file_path.suffix.lower() == '.html':
                    with open(file_path, "r", encoding="utf-8") as f:
                        html_content = f.read()
                    text = self._extract_text_from_html(html_content)
                
                else:
                    continue
                
                # Skip empty documents
                if not text.strip():
                    logger.warning(f"Empty document: {file_path}")
                    continue
                
                # Create document entry with metadata
                doc = {
                    "text": text,
                    "source": str(file_path),
                    "filename": file_path.name,
                    "created_at": os.path.getctime(file_path),
                    "file_type": file_path.suffix.lower()
                }
                documents.append(doc)
                
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")
        
        logger.info(f"Loaded {len(documents)} documents")
        return documents
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from a PDF file using both PyMuPDF and pdfplumber
        for better extraction quality.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text
        """
        text = ""
        
        # Try PyMuPDF first
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            doc.close()
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed for {file_path}: {e}")
            
            # Fallback to pdfplumber
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
            except Exception as e2:
                logger.error(f"Both PDF extraction methods failed for {file_path}: {e2}")
        
        return text
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text
        """
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return ""
    
    def _extract_text_from_html(self, html_content: str) -> str:
        """
        Extract text from HTML content.
        
        Args:
            html_content: HTML content as string
            
        Returns:
            Extracted text
        """
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            # Get text
            text = soup.get_text()
            
            # Break into lines and remove leading and trailing space
            lines = (line.strip() for line in text.splitlines())
            
            # Break multi-headlines into a line each
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            
            # Remove blank lines
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            return ""
    
    def identify_sections(self, text: str) -> List[Dict]:
        """
        Identify sections within a document.
        
        Args:
            text: Document text
            
        Returns:
            List of section dictionaries with text and metadata
        """
        lines = text.split('\n')
        sections = []
        current_section = {"title": "Introduction", "content": "", "start_line": 0}
        
        for i, line in enumerate(lines):
            # Check if this line is a section marker
            is_section_marker = False
            for marker in self.section_markers:
                if re.match(marker, line.strip()):
                    is_section_marker = True
                    # Save previous section if not empty
                    if current_section["content"].strip():
                        sections.append(current_section)
                    
                    # Start new section
                    current_section = {
                        "title": line.strip(),
                        "content": "",
                        "start_line": i
                    }
                    break
            
            # If not a section marker, add to current section content
            if not is_section_marker:
                current_section["content"] += line + "\n"
        
        # Add the last section
        if current_section["content"].strip():
            sections.append(current_section)
        
        return sections
    
    def preprocess_documents(self, documents: List[Dict]) -> List[Dict]:
        """
        Preprocess documents by splitting them into chunks.
        
        Args:
            documents: List of document dictionaries
            
        Returns:
            List of document chunk dictionaries
        """
        processed_docs = []
        
        for doc_idx, doc in enumerate(tqdm(documents, desc="Processing documents")):
            text = doc["text"]
            
            # Identify sections
            sections = self.identify_sections(text)
            
            # Process each section
            for section_idx, section in enumerate(sections):
                section_text = section["content"]
                section_title = section["title"]
                
                # Split section into chunks
                chunks = self.text_splitter.split_text(section_text)
                
                # Create document chunk entries
                for chunk_idx, chunk in enumerate(chunks):
                    chunk_doc = {
                        "text": chunk,
                        "source": doc["source"],
                        "filename": doc["filename"],
                        "section": section_title,
                        "section_idx": section_idx,
                        "chunk_idx": chunk_idx,
                        "doc_idx": doc_idx,
                        "full_reference": f"{doc['filename']} - Section: {section_title}"
                    }
                    processed_docs.append(chunk_doc)
        
        logger.info(f"Created {len(processed_docs)} document chunks")
        self.document_store = processed_docs
        return processed_docs
    
    def create_embeddings(self, documents: Optional[List[Dict]] = None) -> np.ndarray:
        """
        Create embeddings for all documents or provided documents.
        
        Args:
            documents: Optional list of documents to embed
            
        Returns:
            Numpy array of document embeddings
        """
        docs_to_embed = documents if documents is not None else self.document_store
        
        if not docs_to_embed:
            logger.warning("No documents to embed")
            return np.array([])
        
        texts = [doc["text"] for doc in docs_to_embed]
        
        # Process in batches to avoid memory issues with large document collections
        batch_size = 32
        embeddings_list = []
        
        for i in tqdm(range(0, len(texts), batch_size), desc="Creating embeddings"):
            batch_texts = texts[i:i+batch_size]
            batch_embeddings = self.embedding_model.encode(batch_texts)
            embeddings_list.append(batch_embeddings)
        
        embeddings = np.vstack(embeddings_list) if embeddings_list else np.array([])
        
        self.document_embeddings = embeddings
        logger.info(f"Created embeddings with shape {embeddings.shape}")
        return embeddings
    
    def build_index(self, embeddings: Optional[np.ndarray] = None) -> faiss.Index:
        """
        Build a FAISS index for fast similarity search.
        
        Args:
            embeddings: Optional pre-computed embeddings
            
        Returns:
            FAISS index
        """
        embs = embeddings if embeddings is not None else self.document_embeddings
        
        if embs is None or len(embs) == 0:
            raise ValueError("No embeddings available. Call create_embeddings first.")
        
        # Get dimensions of the embeddings
        d = embs.shape[1]
        
        # Create a Flat L2 index (exact search)
        index = faiss.IndexFlatL2(d)
        
        # Add embeddings to the index
        index.add(embs.astype('float32'))
        
        self.index = index
        logger.info(f"Built index with {index.ntotal} vectors")
        return index
    
    def extract_keywords(self, text: str, max_keywords: int = 10) -> Set[str]:
        """
        Extract keywords from text.
        
        Args:
            text: Text to extract keywords from
            max_keywords: Maximum number of keywords to extract
            
        Returns:
            Set of keywords
        """
        doc = self.nlp(text)
        keywords = set()
        
        # Extract named entities
        for ent in doc.ents:
            keywords.add(ent.text.lower())
        
        # Extract noun chunks
        for chunk in doc.noun_chunks:
            keywords.add(chunk.text.lower())
        
        # Get most frequent lemmas that are nouns, verbs, or adjectives
        word_freq = {}
        for token in doc:
            if token.pos_ in ["NOUN", "VERB", "ADJ"] and not token.is_stop and token.lemma_.strip():
                word = token.lemma_.lower()
                if word in word_freq:
                    word_freq[word] += 1
                else:
                    word_freq[word] = 1
        
        # Sort by frequency and add top words
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        for word, _ in sorted_words[:max_keywords]:
            keywords.add(word)
        
        return keywords
    
    def add_keywords_to_documents(self):
        """Add extracted keywords to document chunks."""
        for i, doc in enumerate(tqdm(self.document_store, desc="Extracting keywords")):
            keywords = self.extract_keywords(doc["text"])
            self.document_store[i]["keywords"] = list(keywords)
    
    def process_pipeline(self, docs_dir: Union[str, Path]) -> Tuple[List[Dict], np.ndarray, faiss.Index]:
        """
        Run the full document processing pipeline.
        
        Args:
            docs_dir: Directory containing document files
            
        Returns:
            Tuple of (document_store, embeddings, index)
        """
        documents = self.load_documents(docs_dir)
        processed_docs = self.preprocess_documents(documents)
        self.add_keywords_to_documents()
        embeddings = self.create_embeddings(processed_docs)
        index = self.build_index(embeddings)
        
        return processed_docs, embeddings, index
    
    def save_processed_data(self, save_dir: Union[str, Path]):
        """
        Save processed documents and embeddings.
        
        Args:
            save_dir: Directory to save the processed data
        """
        save_dir = Path(save_dir)
        save_dir.mkdir(parents=True, exist_ok=True)
        
        # Save document store as JSON
        docs_df = pd.DataFrame(self.document_store)
        docs_df.to_json(save_dir / "document_store.json", orient="records")
        
        # Save embeddings as numpy array
        if self.document_embeddings is not None:
            np.save(save_dir / "document_embeddings.npy", self.document_embeddings)
        
        # Save FAISS index if available
        if self.index is not None:
            faiss.write_index(self.index, str(save_dir / "faiss_index.bin"))
        
        logger.info(f"Saved processed data to {save_dir}")
    
    def load_processed_data(self, load_dir: Union[str, Path]):
        """
        Load previously processed documents and embeddings.
        
        Args:
            load_dir: Directory from which to load the processed data
        """
        load_dir = Path(load_dir)
        
        # Load document store
        docs_path = load_dir / "document_store.json"
        if docs_path.exists():
            docs_df = pd.read_json(docs_path)
            self.document_store = docs_df.to_dict(orient="records")
            logger.info(f"Loaded {len(self.document_store)} document chunks")
        
        # Load embeddings
        embs_path = load_dir / "document_embeddings.npy"
        if embs_path.exists():
            self.document_embeddings = np.load(embs_path)
            logger.info(f"Loaded embeddings with shape {self.document_embeddings.shape}")
        
        # Load FAISS index
        index_path = load_dir / "faiss_index.bin"
        if index_path.exists():
            self.index = faiss.read_index(str(index_path))
            logger.info(f"Loaded index with {self.index.ntotal} vectors")


if __name__ == "__main__":
    # Example usage
    processor = DocumentProcessor()
    processor.process_pipeline("data/documents")
    processor.save_processed_data("data/processed")
